# main.py
"""
Agentic RAG Tool - FastAPI Backend
Multi-agent document generation system with Langchain integration
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse, Response, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import json
import asyncio
from datetime import datetime
import uuid
import os
import aiofiles
import logging
import tempfile
from pathlib import Path
from sqlalchemy.orm import Session
from io import BytesIO

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import re

# Import markdown2, but delay weasyprint import until needed
import markdown2

def _create_temp_file(content, extension):
    """Create a temporary file with the given content and extension."""
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{extension}")
    temp_file.write(content.encode('utf-8'))
    temp_file.close()
    return temp_file.name

def _safe_filename(title, extension):
    """Create a safe filename from the title and extension."""
    # Replace invalid characters with underscores
    safe_title = re.sub(r'[^\w\-\.]', '_', title)
    # Limit length and ensure it's not empty
    safe_title = safe_title[:50] or "document"
    return f"{safe_title}.{extension}"

def _markdown_to_latex(markdown_content, title):
    """Simple conversion from Markdown to LaTeX without external dependencies."""
    # Create LaTeX document structure
    latex = [
        "\\documentclass{article}",
        "\\usepackage[utf8]{inputenc}",
        "\\usepackage{hyperref}",
        "\\usepackage{graphicx}",
        "\\usepackage{listings}",
        "\\usepackage{color}",
        "\\usepackage{enumitem}",
        "\\usepackage{booktabs}",
        "\\usepackage{textcomp}",
        "\\DeclareUnicodeCharacter{2013}{--}",  # en-dash
        "\\DeclareUnicodeCharacter{2014}{---}", # em-dash
        "",
        "\\title{" + title.replace("_", "\\_") + "}",
        "\\author{}",
        "\\date{\\today}",
        "",
        "\\begin{document}",
        "",
        "\\maketitle",
        ""
    ]
    
    # Process markdown content line by line
    lines = markdown_content.split('\n')
    in_code_block = False
    in_list = False
    in_table = False
    table_header = []
    
    for line in lines:
        # Handle headings
        if line.startswith('# '):
            latex.append("\\section{" + line[2:].replace("_", "\\_") + "}")
        elif line.startswith('## '):
            latex.append("\\subsection{" + line[3:].replace("_", "\\_") + "}")
        elif line.startswith('### '):
            latex.append("\\subsubsection{" + line[4:].replace("_", "\\_") + "}")
        
        # Handle code blocks
        elif line.startswith('```'):
            if in_code_block:
                latex.append("\\end{lstlisting}")
                in_code_block = False
            else:
                latex.append("\\begin{lstlisting}")
                in_code_block = True
        
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            if not in_list:
                latex.append("\\begin{itemize}")
                in_list = True
            latex.append("  \\item " + line[2:].replace("_", "\\_"))
        elif line.startswith('1. ') or line.startswith('1) '):
            if not in_list:
                latex.append("\\begin{enumerate}")
                in_list = True
            latex.append("  \\item " + line[3:].replace("_", "\\_"))
        elif in_list and line.strip() == '':
            latex.append("\\end{itemize}" if line.startswith('- ') or line.startswith('* ') else "\\end{enumerate}")
            in_list = False
        
        # Handle tables (basic support)
        elif line.startswith('|') and line.endswith('|'):
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            if not in_table:
                in_table = True
                table_header = cells
                latex.append("\\begin{table}[h!]")
                latex.append("\\centering")
                latex.append("\\begin{tabular}{" + "l" * len(cells) + "}")
                latex.append("\\toprule")
                latex.append(" & ".join([cell.replace("_", "\\_") for cell in cells]) + " \\\\")
                latex.append("\\midrule")
            elif all(c.startswith('-') for c in cells):
                # This is the separator row, skip it
                pass
            else:
                latex.append(" & ".join([cell.replace("_", "\\_") for cell in cells]) + " \\\\")
        elif in_table and line.strip() == '':
            latex.append("\\bottomrule")
            latex.append("\\end{tabular}")
            latex.append("\\end{table}")
            in_table = False
        
        # Handle regular text (if not in a special block)
        elif not in_code_block and not line.strip() == '':
            # Basic formatting
            formatted_line = line.replace("**", "\\textbf{").replace("**", "}")
            formatted_line = formatted_line.replace("*", "\\textit{").replace("*", "}")
            formatted_line = formatted_line.replace("_", "\\_")
            # Properly handle dashes
            formatted_line = formatted_line.replace(" - ", " -- ")
            formatted_line = formatted_line.replace("-", "\\textendash{}")
            formatted_line = formatted_line.replace("--", "\\textendash{}")
            formatted_line = formatted_line.replace("---", "\\textemdash{}")
            latex.append(formatted_line)
        
        # Add empty lines
        elif line.strip() == '' and not in_code_block and not in_list and not in_table:
            latex.append("")
    
    # Close any open environments
    if in_code_block:
        latex.append("\\end{lstlisting}")
    if in_list:
        latex.append("\\end{itemize}" if line.startswith('- ') or line.startswith('* ') else "\\end{enumerate}")
    if in_table:
        latex.append("\\bottomrule")
        latex.append("\\end{tabular}")
        latex.append("\\end{table}")
    
    # End document
    latex.append("\\end{document}")
    
    return "\n".join(latex)

from .database import get_db, create_tables, check_database_health
from .models import Document as DBDocument, DocumentChunk, GenerationHistory, UserFeedback
from .workflow import workflow_manager
from .agents.DocumentIngestionAgent import DocumentIngestionAgent
from .agents.ReviewEditingAgent import ReviewEditingAgent
from .config import settings
from .agents.base_agent import safe_serialize_for_db

## NOTE: App initialization moved earlier to ensure decorators work

# FastAPI App
app = FastAPI(
    title="Agentic RAG Tool API",
    description="Multi-agent document generation system with Langchain and Langgraph",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174"],  # Allow both Vite dev server ports
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize database and validate
create_tables()
if not check_database_health():
    logger.error("⚠️ Database schema validation failed. Consider setting DEV_MODE=true to recreate tables or run migrations.")

# Models
class DocumentGenerationRequest(BaseModel):
    doc_type: str
    summary: str
    requirements: str
    style: str = "professional"
    feedback_score: Optional[int] = 3  # RLHF feedback score (1-5)

class DocumentResponse(BaseModel):
    id: str
    content: str
    title: str
    doc_type: str
    status: str
    created_at: str
    updated_at: str
    feedback_score: int  # RLHF feedback score

class ReviewRequest(BaseModel):
    document_id: Optional[str] = None
    doc_id: Optional[str] = None  # Alternative field name for frontend compatibility
    content: str
    feedback: Optional[List[str]] = None
    feedback_score: Optional[int] = 3  # RLHF feedback score

class FeedbackRequest(BaseModel):
    score: int
    feedback: Optional[List[str]] = None

class DocumentListResponse(BaseModel):
    id: str
    title: str
    doc_type: str
    created: str
    modified: str
    status: str
    size: str
    author: str
    feedback_score: float  # RLHF feedback score can be fractional

class UpdateDocumentRequest(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    status: Optional[str] = None
    feedback_score: Optional[int] = None  # RLHF feedback score


@app.options("/api/review")
async def options_review():
    """Handle OPTIONS requests for the review endpoint."""
    return {}

@app.options("/api/review/{doc_id}")
async def options_review_doc():
    """Handle OPTIONS requests for the review document endpoint."""
    return {}

@app.post("/api/review/{doc_id}")
async def review_document_by_id(doc_id: str, request: ReviewRequest = Body(...), db: Session = Depends(get_db)):
    """Apply AI-based formatting improvements and optional feedback-driven edits to a specific document."""
    try:
        # Load document by ID
        document = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        content = request.content or document.content
        if not content:
            raise HTTPException(status_code=400, detail="No content provided for review")

        review_agent = ReviewEditingAgent()
        review_type = "both" if (request.feedback and len(request.feedback) > 0) else "formatting"
        result = await review_agent.execute(
            content=content,
            doc_type=document.doc_type,
            style_profile=document.style_metadata if document.style_metadata else {},
            feedback=request.feedback or [],
            review_type=review_type,
            approved=False,
            feedback_score=request.feedback_score or 3,
            db_session=None,
        )

        if result.get("status") != "success":
            raise HTTPException(status_code=500, detail=result.get("message", "Review failed"))

        # Update document with improved content
        document.content = result.get("improved_content", content)
        document.updated_at = datetime.utcnow()
        db.commit()

        # Return fields the frontend can consume (supports both snake_case and camelCase)
        return {
            "status": "success",
            "document_id": doc_id,
            "improved_content": result.get("improved_content", content),
            "improvedContent": result.get("improved_content", content),
            "changes_made": result.get("changes_made", []),
            "changesMade": result.get("changes_made", []),
            "suggestions": [],
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Review endpoint failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/review")
async def review_document(request: ReviewRequest, db: Session = Depends(get_db)):
    """Apply AI-based formatting improvements and optional feedback-driven edits."""
    try:
        # Load document if ID provided; otherwise operate on provided content only
        document = None
        doc_id = request.document_id or request.doc_id
        
        if doc_id:
            document = db.query(DBDocument).filter(DBDocument.id == doc_id).first()
            if not document:
                raise HTTPException(status_code=404, detail="Document not found")

        content = request.content or (document.content if document else "")
        if not content:
            raise HTTPException(status_code=400, detail="No content provided for review")

        review_agent = ReviewEditingAgent()
        review_type = "both" if (request.feedback and len(request.feedback) > 0) else "formatting"
        result = await review_agent.execute(
            content=content,
            doc_type=document.doc_type if document else "SRS",
            style_profile=document.style_metadata if document and document.style_metadata else {},
            feedback=request.feedback or [],
            review_type=review_type,
            approved=False,
            feedback_score=request.feedback_score or 3,
            db_session=None,
        )

        if result.get("status") != "success":
            raise HTTPException(status_code=500, detail=result.get("message", "Review failed"))

        # Return fields the frontend can consume (supports both snake_case and camelCase)
        return {
            "status": "success",
            "doc_id": doc_id,
            "document_id": doc_id,
            "updated_content": result.get("improved_content", content),
            "improved_content": result.get("improved_content", content),
            "improvedContent": result.get("improved_content", content),
            "changes_made": result.get("changes_made", []),
            "changesMade": result.get("changes_made", []),
            "suggestions": [],
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Review endpoint failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Note: FastAPI App is already initialized above, no need to reinitialize

# Add alias routes for /api/docs endpoints
@app.get("/api/docs")
async def get_all_documents(db: Session = Depends(get_db)):
    """Get all documents - alias for /api/documents."""
    return await list_documents(db)

@app.get("/api/docs/{doc_id}")
async def get_document_by_id(doc_id: str, db: Session = Depends(get_db)):
    """Get a document by ID - alias for /api/documents/{document_id}."""
    return await get_document(doc_id, db)

# Initialize database and validate
create_tables()
if not check_database_health():
    logger.error("⚠️ Database schema validation failed. Consider setting DEV_MODE=true to recreate tables or run migrations.")

def _safe_filename(title: str, ext: str) -> str:
    """Create a filesystem-safe filename with the given extension."""
    safe = "".join(c if c.isalnum() or c in (" ", "-", "_") else "_" for c in (title or "document")).strip()
    safe = safe.replace(" ", "_")
    return f"{safe or 'document'}.{ext}"

def _format_markdown(content: str) -> str:
    """Format Markdown with consistent section numbering and table alignment."""
    content = re.sub(r'\n{3,}', '\n\n', content)
    content = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', content)
    lines = content.split('\n')
    processed_lines = []
    section_number = [0]
    
    def update_section_number(line, level):
        if level == 1:
            section_number[0] += 1
            section_number[1:] = [0] * (len(section_number) - 1)
        elif level <= len(section_number):
            section_number[level - 1] += 1
            section_number[level:] = [0] * (len(section_number) - level)
        else:
            section_number.extend([0] * (level - len(section_number) - 1))
            section_number.append(1)
        return '.'.join(str(n) for n in section_number[:level]) + ' ' + line.lstrip('#').strip()
    
    for line in lines:
        if line.startswith('#'):
            level = line.count('#')
            processed_lines.append('')
            processed_lines.append(f"{'#' * level} {update_section_number(line, level)}")
            processed_lines.append('')
        elif line.startswith('|'):
            processed_lines.append(line.replace(' | ', '|').replace('| ', '|').replace(' |', '|'))
        else:
            processed_lines.append(line)
    
    final_lines = []
    blank_count = 0
    for line in processed_lines:
        if not line.strip():
            blank_count += 1
            if blank_count <= 2:
                final_lines.append(line)
        else:
            blank_count = 0
            final_lines.append(line)
    
    return '\n'.join(final_lines).strip()

@app.get("/")
async def root():
    return {
        "message": "Agentic RAG Tool API", 
        "status": "active",
        "version": "1.0.0",
        "agents": list(workflow_manager.agents.keys()),
        "db_health": check_database_health()
    }

@app.options("/api/upload")
async def options_upload():
    """Handle OPTIONS requests for the upload endpoint."""
    return {}

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Upload and process a document using DocumentIngestionAgent."""
    try:
        # Validate file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in settings.allowed_extensions:
            raise HTTPException(status_code=400, detail=f"File type {file_ext} not supported. Allowed: {settings.allowed_extensions}")

        # Validate file size
        content = await file.read()
        if len(content) > settings.max_file_size:
            raise HTTPException(status_code=400, detail=f"File size exceeds limit of {settings.max_file_size} bytes")

        # Create temp file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_path = Path(temp_file.name)
            
        # Use aiofiles correctly with async with
        async with aiofiles.open(temp_path, mode='wb') as f:
            await f.write(content)

        # Process with ingestion agent
        agent = DocumentIngestionAgent()
        result = await agent.execute(
            db=db,
            filename=file.filename,
            content=content.decode('utf-8', errors='ignore'),
            doc_type=file_ext.lstrip('.').upper(),
            approved=True,
            feedback_score=3
        )

        if result["status"] != "success":
            raise HTTPException(status_code=500, detail=result.get("message", "Ingestion failed"))

        # Cleanup temp file
        temp_path.unlink(missing_ok=True)

        return {
            "status": "success",
            "document_id": result["document_id"],
            "message": "Document uploaded and processed successfully"
        }

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.options("/api/generate")
async def options_generate():
    """Handle OPTIONS requests for the generate endpoint."""
    return {}

@app.post("/api/generate")
async def generate_document(request: DocumentGenerationRequest, db: Session = Depends(get_db)):
    """Generate a new document using the workflow manager."""
    try:
        result = await workflow_manager.execute_generation_workflow(
            doc_type=request.doc_type,
            summary=request.summary,
            requirements=request.requirements,
            style=request.style,
            db=db,
            max_iterations=3
        )
        
        if result["status"] == "failed":
            raise HTTPException(status_code=500, detail=result["error"])
        
        # Create document record
        document = DBDocument(
            title=f"{request.doc_type}: {request.summary}",
            filename=f"{request.doc_type}_{uuid.uuid4().hex}.md",
            doc_type=request.doc_type,
            content=result["content"],
            status="final",
            approved=True,
            feedback_score=max(1, min(5, request.feedback_score or 3))
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        return DocumentResponse(
            id=document.id,
            content=document.content,
            title=document.title,
            doc_type=document.doc_type,
            status=document.status,
            created_at=document.created_at.isoformat(),
            updated_at=document.updated_at.isoformat(),
            feedback_score=document.feedback_score
        )
        
    except Exception as e:
        logger.error(f"Generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@app.options("/api/documents")
async def options_list_documents():
    """Handle OPTIONS requests for the documents list endpoint."""
    return {}

@app.get("/api/documents")
async def list_documents(db: Session = Depends(get_db)):
    """List all documents with metadata."""
    try:
        documents = db.query(DBDocument).all()
        
        response = []
        for doc in documents:
            response.append(DocumentListResponse(
                id=doc.id,
                title=doc.title,
                doc_type=doc.doc_type,
                created=doc.created_at.isoformat(),
                modified=doc.updated_at.isoformat(),
                status=doc.status,
                size=f"{len(doc.content.split()) if doc.content else 0} words",
                author="System",
                feedback_score=doc.feedback_score
            ))
        
        return response
        
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")

@app.get("/api/documents/{document_id}")
async def get_document(document_id: str, db: Session = Depends(get_db)):
    """Get a specific document by ID."""
    try:
        document = db.query(DBDocument).filter(DBDocument.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return DocumentResponse(
            id=document.id,
            content=document.content,
            title=document.title,
            doc_type=document.doc_type,
            status=document.status,
            created_at=document.created_at.isoformat(),
            updated_at=document.updated_at.isoformat(),
            feedback_score=document.feedback_score
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document: {str(e)}")

@app.put("/api/documents/{document_id}")
async def update_document(document_id: str, request: UpdateDocumentRequest, db: Session = Depends(get_db)):
    """Update a document's metadata."""
    try:
        document = db.query(DBDocument).filter(DBDocument.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        if request.title is not None:
            document.title = request.title
        
        if request.content is not None:
            document.content = request.content
        
        if request.status is not None:
            document.status = request.status
        
        if request.feedback_score is not None:
            document.feedback_score = max(1, min(5, request.feedback_score))
        
        document.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(document)
        
        return DocumentResponse(
            id=document.id,
            content=document.content,
            title=document.title,
            doc_type=document.doc_type,
            status=document.status,
            created_at=document.created_at.isoformat(),
            updated_at=document.updated_at.isoformat(),
            feedback_score=document.feedback_score
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update document: {str(e)}")

@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: str, db: Session = Depends(get_db)):
    """Delete a document by ID."""
    try:
        document = db.query(DBDocument).filter(DBDocument.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        db.delete(document)
        db.commit()
        
        return {"status": "success", "message": "Document deleted"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@app.get("/api/export/{document_id}")
async def export_document(document_id: str, format: str = "md", db: Session = Depends(get_db)):
    """Export a document in the specified format (md, pdf, docx, latex)."""
    try:
        document = db.query(DBDocument).filter(DBDocument.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        content = document.content
        title = document.title or "document"
        ext = format.lower()
        
        if ext == "md":
            # Return as Markdown
            return FileResponse(
                path=_create_temp_file(content, "md"),
                media_type="text/markdown",
                filename=_safe_filename(title, "md")
            )
        
        elif ext == "pdf":
            try:
                # Try xhtml2pdf first for better table support
                try:
                    from xhtml2pdf import pisa
                    
                    # Pre-process content to ensure proper dash rendering
                    processed_content = content
                    # Ensure proper dash rendering by using HTML entities
                    processed_content = processed_content.replace(" - ", " – ")
                    processed_content = processed_content.replace("--", "—")
                    
                    # Convert Markdown to HTML with proper symbol handling
                    html_content = markdown2.markdown(processed_content, extras=["tables", "fenced-code-blocks"])
                    html = f"""
                    <html>
                    <head>
                        <title>{title}</title>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 20px; }}
                            h1 {{ color: #333366; }}
                            h2 {{ color: #333366; }}
                            table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
                            th, td {{ border: 1px solid #ddd; padding: 8px; }}
                            th {{ background-color: #f2f2f2; }}
                            code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 4px; }}
                            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; overflow-x: auto; }}
                        </style>
                    </head>
                    <body>
                        {html_content}
                    </body>
                    </html>
                    """
                    
                    # Create PDF
                    pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                    pdf_file.close()
                    
                    with open(pdf_file.name, "wb") as output_file:
                        pisa_status = pisa.CreatePDF(html, dest=output_file)
                    
                    if pisa_status.err:
                        raise Exception("PDF generation failed")
                    
                    return FileResponse(
                        path=pdf_file.name,
                        media_type="application/pdf",
                        filename=_safe_filename(title, "pdf")
                    )
                    
                except ImportError:
                    # Fall back to weasyprint if xhtml2pdf is not available
                    from weasyprint import HTML as WeasyHTML
                    
                    # Pre-process content to ensure proper dash rendering
                    processed_content = content
                    # Ensure proper dash rendering by using Unicode characters
                    processed_content = processed_content.replace(" - ", " – ")
                    processed_content = processed_content.replace("--", "—")
                    
                    # Convert Markdown to HTML with proper symbol handling
                    html_content = markdown2.markdown(processed_content, extras=["tables", "fenced-code-blocks"])
                    html = f"""
                    <html>
                    <head>
                        <title>{title}</title>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 20px; }}
                            h1 {{ color: #333366; }}
                            h2 {{ color: #333366; }}
                            table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
                            th, td {{ border: 1px solid #ddd; padding: 8px; }}
                            th {{ background-color: #f2f2f2; }}
                            code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 4px; }}
                            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; overflow-x: auto; }}
                        </style>
                    </head>
                    <body>
                        {html_content}
                    </body>
                    </html>
                    """
                    
                    # Create PDF
                    pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                    pdf_file.close()
                    
                    WeasyHTML(string=html).write_pdf(pdf_file.name)
                    
                    return FileResponse(
                        path=pdf_file.name,
                        media_type="application/pdf",
                        filename=_safe_filename(title, "pdf")
                    )
                    
            except ImportError as e:
                logger.error(f"PDF export dependencies missing: {e}")
                raise HTTPException(status_code=501, detail="PDF export not available due to missing dependencies. Please install xhtml2pdf or weasyprint.")
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")
                raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")
        
        elif ext == "docx":
            try:
                import pypandoc
                
                # Create a temporary markdown file
                md_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md")
                
                # Pre-process content to ensure proper dash rendering
                processed_content = content
                # Ensure proper dash rendering by using HTML entities
                processed_content = processed_content.replace(" - ", " – ")
                processed_content = processed_content.replace("--", "—")
                
                md_file.write(processed_content.encode('utf-8'))
                md_file.close()
                
                # Create output docx file
                docx_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                docx_file.close()
                
                # Convert using pandoc
                try:
                    # Use extra_args to ensure proper Unicode handling
                    pypandoc.convert_file(md_file.name, 'docx', outputfile=docx_file.name, 
                                         extra_args=['--ascii'])
                    os.unlink(md_file.name)  # Clean up the temp markdown file
                    
                    return FileResponse(
                        path=docx_file.name,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=_safe_filename(title, "docx")
                    )
                except Exception as e:
                    # If pypandoc fails, try using python-docx directly
                    from docx import Document
                    
                    doc = Document()
                    doc.add_heading(title, 0)
                    
                    # Enhanced markdown parsing for docx with proper symbol handling
                    paragraphs = processed_content.split('\n\n')
                    for para in paragraphs:
                        if para.startswith('# '):
                            doc.add_heading(para[2:], 1)
                        elif para.startswith('## '):
                            doc.add_heading(para[3:], 2)
                        elif para.startswith('### '):
                            doc.add_heading(para[4:], 3)
                        elif para.startswith('- ') or para.startswith('* '):
                            # Handle list items
                            p = doc.add_paragraph()
                            p.style = 'List Bullet'
                            p.add_run(para[2:])
                        else:
                            doc.add_paragraph(para)
                    
                    doc.save(docx_file.name)
                    
                    return FileResponse(
                        path=docx_file.name,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=_safe_filename(title, "docx")
                    )
                    
            except ImportError as e:
                logger.error(f"DOCX export dependencies missing: {e}")
                raise HTTPException(status_code=501, detail="DOCX export not available due to missing dependencies. Please install pypandoc or python-docx.")
        
        elif ext == "latex":
            try:
                # First try with pypandoc
                try:
                    import pypandoc
                    
                    # Create a temporary markdown file
                    md_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md")
                    md_file.write(content.encode('utf-8'))
                    md_file.close()
                    
                    # Create output latex file
                    latex_file = tempfile.NamedTemporaryFile(delete=False, suffix=".tex")
                    latex_file.close()
                    
                    # Convert using pandoc
                    pypandoc.convert_file(md_file.name, 'latex', outputfile=latex_file.name)
                    os.unlink(md_file.name)  # Clean up the temp markdown file
                    
                    # Read the LaTeX content
                    with open(latex_file.name, 'r', encoding='utf-8') as f:
                        latex_content = f.read()
                    
                    return FileResponse(
                        path=latex_file.name,
                        media_type="application/x-latex",
                        filename=_safe_filename(title, "tex")
                    )
                    
                except (ImportError, OSError, Exception) as e:
                    logger.warning(f"Pypandoc conversion failed: {e}. Using manual conversion.")
                    # If pypandoc fails, use a simple manual conversion
                    latex_content = _markdown_to_latex(content, title)
                    
                    # Create a temporary LaTeX file
                    latex_file = tempfile.NamedTemporaryFile(delete=False, suffix=".tex")
                    latex_file.write(latex_content.encode('utf-8'))
                    latex_file.close()
                    
                    return FileResponse(
                        path=latex_file.name,
                        media_type="application/x-latex",
                        filename=_safe_filename(title, "tex")
                    )
                
            except Exception as e:
                logger.error(f"LaTeX export failed: {e}")
                raise HTTPException(status_code=500, detail=f"LaTeX export failed: {str(e)}")
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported format. Supported: md, pdf, docx, latex")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)