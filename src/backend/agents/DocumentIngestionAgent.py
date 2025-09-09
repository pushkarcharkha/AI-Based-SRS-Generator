"""Document Ingestion Agent - Production-ready implementation for uploading and processing docs"""

import os
import aiofiles
from typing import Dict, Any, List, Optional
from sqlalchemy.orm import Session
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
import uuid
import re
import json
import asyncio
import tenacity
import logging

from .base_agent import BaseAgent
from ..models import Document, DocumentChunk
from ..vector_store import VectorStoreWrapper
from ..config import settings

logger = logging.getLogger(__name__)

class DocumentIngestionAgent(BaseAgent):
    """Production-ready Document Ingestion Agent for processing and storing documents"""
    
    def __init__(self):
        super().__init__(name="document_ingestion", description="Supports uploading docs, extracting style + content, and exporting in multiple formats")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len
        )
        self.vector_store = VectorStoreWrapper()
    
    async def _parse_document(self, file_path: str, filename: str) -> str:
        """Parse document content based on file extension"""
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            if file_ext == '.pdf':
                import pypdf
                from io import BytesIO
                
                async with aiofiles.open(file_path, 'rb') as file:
                    content = await file.read()
                
                pdf_reader = pypdf.PdfReader(BytesIO(content))
                text_parts = []
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {i+1}]\n{page_text}")
                
                return "\n\n".join(text_parts)
            
            elif file_ext == '.docx':
                import docx
                from io import BytesIO
                
                async with aiofiles.open(file_path, 'rb') as file:
                    content = await file.read()
                
                doc = docx.Document(BytesIO(content))
                text_parts = []
                
                # Extract paragraphs
                for p in doc.paragraphs:
                    if p.text.strip():
                        text_parts.append(p.text)
                
                # Extract table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                return "\n\n".join(text_parts)
            
            elif file_ext in ['.txt', '.md']:
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = await file.read()
                return content.strip()
            
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        
        except Exception as e:
            logger.error(f"Error parsing {filename}: {e}")
            raise ValueError(f"Failed to parse document: {str(e)}")
    
    def _extract_style_metadata(self, content: str) -> Dict[str, Any]:
        """Extract style patterns from content for formatting preservation"""
        try:
            metadata = {
                "heading_patterns": {
                    "hash_headers": len(re.findall(r'^#+\s', content, re.M)),
                    "numbered_sections": len(re.findall(r'^\d+\.\s', content, re.M)),
                    "lettered_sections": len(re.findall(r'^[a-zA-Z]\.\s', content, re.M))
                },
                "list_indicators": {
                    "bullet_points": len(re.findall(r'^[-*•]\s', content, re.M)),
                    "numbered_lists": len(re.findall(r'^\d+\)\s', content, re.M)),
                    "dash_lists": len(re.findall(r'^-\s', content, re.M))
                },
                "formatting_patterns": {
                    "bold_text": len(re.findall(r'\*\*.*?\*\*', content)),
                    "italic_text": len(re.findall(r'\*.*?\*', content)),
                    "code_blocks": len(re.findall(r'```.*?```', content, re.DOTALL)),
                    "inline_code": len(re.findall(r'`.*?`', content))
                }
            }
            return metadata
        except Exception as e:
            logger.warning(f"Error extracting style metadata: {e}")
            return {}
    
    def _detect_document_type(self, content: str, filename: str) -> str:
        """Detect document type based on content and filename patterns"""
        try:
            content_lower = content.lower()
            filename_lower = filename.lower()
            
            patterns = {
                'SRS': (
                    ['srs', 'requirements', 'specification', 'req'], 
                    ['software requirements', 'functional requirements', 'non-functional requirements']
                ),
                'SOW': (
                    ['sow', 'statement', 'work', 'scope'], 
                    ['statement of work', 'deliverables', 'timeline', 'project scope']
                ),
                'Proposal': (
                    ['proposal', 'rfp', 'bid', 'quote'], 
                    ['proposal', 'budget', 'cost estimate']
                ),
                'Technical': (
                    ['technical', 'api', 'documentation', 'tech', 'guide'], 
                    ['api documentation', 'technical specification', 'architecture']
                ),
                'Business': (
                    ['business', 'plan', 'strategy', 'market'], 
                    ['business plan', 'market analysis', 'financial projections']
                )
            }
            
            for doc_type, (file_patterns, content_patterns) in patterns.items():
                # Check filename patterns
                if any(pattern in filename_lower for pattern in file_patterns):
                    return doc_type
                
                # Check content patterns
                content_score = sum(content_lower.count(pattern) for pattern in content_patterns)
                if content_score > 0:
                    return doc_type
            
            return 'General'
            
        except Exception as e:
            logger.warning(f"Error detecting document type: {e}")
            return 'General'
    
    async def _chunk_content(self, content: str, document_id: str) -> List[Dict[str, Any]]:
        """Split content into chunks with metadata"""
        try:
            text_chunks = self.text_splitter.split_text(content)
            chunks = []
            
            for i, chunk in enumerate(text_chunks):
                if chunk.strip():  # Only include non-empty chunks
                    chunks.append({
                        "content": chunk.strip(),
                        "metadata": {
                            "chunk_index": i,
                            "word_count": len(chunk.split()),
                            "char_count": len(chunk),
                            "document_id": document_id
                        }
                    })
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking content: {e}")
            raise ValueError(f"Failed to chunk content: {str(e)}")
    
    @tenacity.retry(
        stop=tenacity.stop_after_attempt(3),
        wait=tenacity.wait_exponential(multiplier=1, min=4, max=10),
        retry=tenacity.retry_if_exception_type(Exception)
    )
    async def execute(
        self, 
        db: Session, 
        filename: str, 
        file_path: Optional[str] = None, 
        content: Optional[str] = None, 
        doc_type: str = "auto-detect", 
        approved: bool = False, 
        feedback_score: int = 3
    ) -> Dict[str, Any]:
        """Ingest and process a document"""
        try:
            # Parse document content if not provided
            if content is None:
                if not file_path:
                    raise ValueError("Either content or file_path must be provided")
                content = await self._parse_document(file_path, filename)
            
            if not content.strip():
                raise ValueError("No content extracted from document")
            
            # Extract style metadata and detect document type
            style_metadata = self._extract_style_metadata(content)
            if doc_type == "auto-detect":
                doc_type = self._detect_document_type(content, filename)
            
            # Create document record
            document_id = str(uuid.uuid4())
            document = Document(
                id=document_id,
                filename=filename,
                title=filename.replace('_', ' ').replace('-', ' ').title(),
                doc_type=doc_type,
                content=content,
                file_path=file_path,
                file_size=len(content),
                style_metadata=json.dumps(style_metadata),
                status="processing",
                approved=approved,
                feedback_score=max(1, min(5, feedback_score))
            )
            
            # Add and flush document to get ID
            db.add(document)
            db.flush()
            
            # Chunk the content
            chunks = await self._chunk_content(content, document_id)
            db_chunks = []
            langchain_docs = []
            
            for i, chunk_data in enumerate(chunks):
                if not chunk_data["content"].strip():
                    continue
                
                metadata = chunk_data.get("metadata", {})
                metadata.update({
                    "approved": approved, 
                    "feedback_score": feedback_score, 
                    "document_id": document.id
                })
                
                chunk_id = str(uuid.uuid4())
                db_chunk = DocumentChunk(
                    id=chunk_id,
                    document_id=document.id,
                    content=chunk_data["content"],
                    chunk_index=i,
                    metadata=json.dumps(metadata),
                    embedding_model=settings.embedding_model
                )
                db_chunks.append(db_chunk)
                
                # Create Langchain document for vector store
                lc_doc = LCDocument(
                    page_content=chunk_data["content"], 
                    metadata=metadata
                )
                langchain_docs.append(lc_doc)
            
            # Add chunks to database
            if db_chunks:
                db.add_all(db_chunks)
            
            # Add to vector store
            vector_ids = []
            if langchain_docs:
                try:
                    vector_ids = await self.vector_store.async_add_documents(langchain_docs)
                except Exception as e:
                    logger.warning(f"Vector store addition failed: {e}")
                    # Continue without vector IDs if vector store fails
            
            # Update document status
            document.status = "completed"
            
            # Commit all changes
            db.commit()
            
            return {
                "status": "success",
                "document_id": document.id,
                "chunk_count": len(chunks),
                "vector_ids": vector_ids,
                "metadata": {
                    "filename": filename,
                    "doc_type": doc_type,
                    "file_size": document.file_size,
                    "word_count": len(content.split()),
                    "style_metadata": style_metadata,
                    "feedback_score": feedback_score
                }
            }
            
        except Exception as e:
            # Rollback on error
            db.rollback()
            logger.error(f"DocumentIngestionAgent execution failed: {e}", exc_info=True)
            return {
                "status": "error",
                "message": str(e),
                "document_id": None
            }
    
    @tenacity.retry(
        stop=tenacity.stop_after_attempt(3),
        wait=tenacity.wait_exponential(multiplier=1, min=4, max=10),
        retry=tenacity.retry_if_exception_type(Exception)
    )
    async def update_feedback(self, db: Session, document_id: str, score: int) -> Dict[str, Any]:
        """Update feedback score for document and associated chunks"""
        try:
            # Find document
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise ValueError("Document not found")
            
            # Update document feedback score
            document.feedback_score = max(1, min(5, score))
            
            # Update associated chunks
            chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).all()
            for chunk in chunks:
                try:
                    metadata = json.loads(chunk.metadata) if chunk.metadata else {}
                    metadata["feedback_score"] = score
                    chunk.metadata = json.dumps(metadata)
                    
                    # Update vector store metadata if vector_id exists
                    if hasattr(self.vector_store, 'async_update_metadata') and chunk.vector_id:
                        try:
                            await self.vector_store.async_update_metadata(
                                chunk.vector_id, 
                                {"feedback_score": score}
                            )
                        except Exception as e:
                            logger.warning(f"Failed to update vector store metadata: {e}")
                            
                except Exception as e:
                    logger.warning(f"Failed to update chunk {chunk.id} metadata: {e}")
            
            db.commit()
            return {
                "status": "success", 
                "updated_score": score
            }
            
        except Exception as e:
            db.rollback()
            logger.error(f"Feedback update failed: {e}")
            return {
                "status": "error",
                "message": str(e)
            }
    
    @tenacity.retry(
        stop=tenacity.stop_after_attempt(3),
        wait=tenacity.wait_exponential(multiplier=1, min=4, max=10),
        retry=tenacity.retry_if_exception_type(Exception)
    )
    def _create_latex_header_file(self) -> str:
        """Create a temporary file with LaTeX header configurations for Unicode support"""
        import tempfile
        import os
        
        # Create a temporary file for the LaTeX header
        fd, path = tempfile.mkstemp(suffix='.tex')
        
        # LaTeX header content with Unicode configuration
        header_content = r'''
\usepackage{fontspec}
\usepackage{unicode-math}
\usepackage{microtype}
\usepackage{xeCJK}
\usepackage{ucharclasses}
\usepackage{listings}
\usepackage{textcomp}

% Ensure proper dash handling
\DeclareUnicodeCharacter{2013}{--}  % en-dash
\DeclareUnicodeCharacter{2014}{---} % em-dash
\DeclareUnicodeCharacter{2022}{\textbullet} % bullet
\DeclareUnicodeCharacter{2023}{\textbullet} % triangular bullet

% Configure font features for better Unicode support
\defaultfontfeatures{Ligatures=TeX, Scale=MatchLowercase}
'''
        
        # Write the header content to the temporary file
        with os.fdopen(fd, 'w') as f:
            f.write(header_content)
        
        return path
    
    async def export_document(self, db: Session, document_id: str, format: str = "md") -> bytes:
        """Export document in specified format"""
        try:
            # Find document
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise ValueError("Document not found")
            
            # Ensure proper Unicode handling for special characters like dashes
            # Replace problematic dash characters with their proper Unicode equivalents
            processed_content = document.content
            # Replace en-dash and em-dash with their proper Unicode representations
            processed_content = processed_content.replace('–', '\u2013').replace('—', '\u2014')
            # Replace hyphen-minus with proper hyphen
            processed_content = processed_content.replace('-', '\u002D')
            # Replace other potentially problematic characters if needed
            content = processed_content.encode('utf-8')
            
            if format == "md":
                return content
            elif format in ["pdf", "docx", "latex"]:
                # Use pandoc for conversion if available
                try:
                    import subprocess
                    to_format = "latex" if format == "latex" else format
                    
                    # Configure Pandoc to use XeLaTeX with Unicode font support for PDF
                    pandoc_args = ['pandoc', '-f', 'markdown', '-t', to_format]
                    header_file_path = None
                    
                    # For PDF, use XeLaTeX with Unicode font support
                    if format == "pdf":
                        # Use a list of common Unicode-compatible fonts that work across platforms
                        # The first available font in the list will be used
                        header_file_path = self._create_latex_header_file()
                        pandoc_args.extend([
                            '--pdf-engine=xelatex',
                            # Try multiple fonts in order of preference
                            # Windows fonts first, then cross-platform fonts
                            '-V', 'mainfont=Arial, Calibri, Times New Roman, Liberation Serif, DejaVu Sans, Noto Sans, Segoe UI',
                            '-V', 'sansfont=Arial, Calibri, Segoe UI, Liberation Sans, DejaVu Sans, Noto Sans',
                            '-V', 'monofont=Consolas, Courier New, Liberation Mono, DejaVu Sans Mono',
                            '--variable=fontsize:11pt',
                            # Add variables to ensure proper Unicode handling
                            '-V', 'geometry:margin=1in',
                            '-V', 'documentclass=article',
                            # Use the custom LaTeX template with proper font configuration
                            '--include-in-header=' + header_file_path
                        ])
                    
                    try:
                        proc = await asyncio.create_subprocess_exec(
                            *pandoc_args,
                            stdin=asyncio.subprocess.PIPE, 
                            stdout=asyncio.subprocess.PIPE,
                            stderr=asyncio.subprocess.PIPE
                        )
                        
                        stdout, stderr = await proc.communicate(content)
                        return_code = await proc.wait()
                        
                        if return_code != 0:
                            error_msg = stderr.decode('utf-8') if stderr else "Unknown pandoc error"
                            raise RuntimeError(f"Pandoc conversion failed: {error_msg}")
                        
                        return stdout
                    finally:
                        # Clean up the temporary header file if it was created
                        if header_file_path and os.path.exists(header_file_path):
                            try:
                                os.remove(header_file_path)
                                logger.debug(f"Removed temporary LaTeX header file: {header_file_path}")
                            except Exception as e:
                                logger.warning(f"Failed to remove temporary file {header_file_path}: {e}")

                    
                except FileNotFoundError:
                    # If pandoc is not available, return markdown
                    logger.warning("Pandoc not found, returning markdown instead")
                    return content
                except Exception as e:
                    logger.error(f"Pandoc conversion failed: {e}")
                    # Return markdown as fallback
                    return content
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            logger.error(f"Export failed: {e}")
            raise